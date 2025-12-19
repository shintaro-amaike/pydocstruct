"""pydocstruct/loaders/docx_loader.py"""
from __future__ import annotations

from pathlib import Path
from typing import Any

try:
    import docx
except ImportError:
    docx = None

from pydocstruct.core.document import Document
from pydocstruct.core.loader import BaseLoader


class DocxLoader(BaseLoader):
    """DOCXファイルを読み込むローダー"""

    def __init__(
        self,
        file_path: str | Path,
        **kwargs: Any,
    ) -> None:
        super().__init__(file_path, **kwargs)
        
        if docx is None:
            raise ImportError(
                "python-docxがインストールされていません。"
                "pip install python-docx でインストールしてください。"
            )

    def load(self) -> list[Document]:
        """DOCXファイルを読み込む"""
        doc = docx.Document(self.file_path)
        
        # パラグラフからテキストを抽出
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
                
        content = "\n".join(full_text)
        
        metadata = self._create_base_metadata()
        
        # DOCX固有のメタデータ（もしあれば）
        if hasattr(doc, "core_properties"):
            core_props = doc.core_properties
            metadata["docx_metadata"] = {
                "author": core_props.author,
                "created": str(core_props.created),
                "modified": str(core_props.modified),
                "title": core_props.title,
            }

        return [
            Document(
                content=content,
                metadata=metadata,
                source=str(self.file_path),
            )
        ]

